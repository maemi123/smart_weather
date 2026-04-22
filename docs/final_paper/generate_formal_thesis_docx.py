from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.image.image import Image
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"
FINAL_DIR = DOCS / "final_paper"
TEMPLATE_PATH = FINAL_DIR / "school_template_source.docx"
DRAFT_PATH = DOCS / "thesis_draft.md"
METRICS_PATH = ROOT / "models" / "metrics_report.json"
OUTPUT_ASCII = FINAL_DIR / "formal_thesis_draft.docx"
OUTPUT_CN = FINAL_DIR / "毕业论文初稿_正式排版.docx"

TITLE_CN = "基于 Flask 的智能气象信息服务系统设计与实现"
TITLE_EN = "Design and Implementation of an Intelligent Meteorological Information Service System Based on Flask"
SCHOOL = "山西财经大学"
GRAD_YEAR = "2026"
AUTHOR = "徐毅成"
STUDENT_ID = "202207020136"
COLLEGE = "待填写"
MAJOR = "待填写"
CLASS_NAME = "待填写"
ADVISOR = "待填写"

HEADER_TEXT = f"{SCHOOL}{GRAD_YEAR}届本科生毕业论文（设计）"
LOGO_PATH = FINAL_DIR / "template_media_extract" / "image7.png"


def set_run_font(run, font_name: str, size_pt: float | None = None, *, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, *, first_line_chars: int | None = None, line_spacing: float | None = None,
                         before_pt: float | None = None, after_pt: float | None = None,
                         alignment=None) -> None:
    pf = paragraph.paragraph_format
    if first_line_chars is not None:
        pf.first_line_indent = Pt(21 * first_line_chars / 2)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if before_pt is not None:
        pf.space_before = Pt(before_pt)
    if after_pt is not None:
        pf.space_after = Pt(after_pt)
    if alignment is not None:
        paragraph.alignment = alignment


def keep_paragraph_together(paragraph, *, with_next: bool = False) -> None:
    pf = paragraph.paragraph_format
    pf.keep_together = True
    pf.keep_with_next = with_next


def set_paragraph_bottom_border(paragraph, *, size: int = 6, space: int = 1, color: str = "808080") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在 Word 中更新后显示"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def set_page_number_format(section, fmt: str = "decimal", start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def configure_section(section, *, with_header: bool, with_footer_page_number: bool,
                      page_fmt: str | None, page_start: int | None = None) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    odd_header = section.header.paragraphs[0]
    even_header = section.even_page_header.paragraphs[0]
    odd_footer = section.footer.paragraphs[0]
    even_footer = section.even_page_footer.paragraphs[0]

    for para in (odd_header, even_header, odd_footer, even_footer):
        para.clear()

    odd_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    even_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    odd_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    even_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if with_header:
        for para in (odd_header, even_header):
            run = para.add_run(HEADER_TEXT)
            set_run_font(run, "宋体", 9)
            set_paragraph_bottom_border(para)

    if with_footer_page_number:
        for para in (odd_footer, even_footer):
            add_page_number_field(para)

    if page_fmt:
        set_page_number_format(section, fmt=page_fmt, start=page_start)


def strip_markdown_inline(text: str) -> str:
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def parse_thesis_markdown(text: str) -> dict:
    lines = text.splitlines()
    data: dict = {"title": TITLE_CN, "abstract": [], "keywords": "", "blocks": []}
    current = None
    para_buf: list[str] = []

    def flush_paragraph():
        nonlocal para_buf
        if para_buf and current == "body":
            data["blocks"].append({"type": "paragraph", "text": strip_markdown_inline(" ".join(para_buf))})
        elif para_buf and current == "abstract":
            data["abstract"].append(strip_markdown_inline(" ".join(para_buf)))
        para_buf = []

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        raw = line.rstrip()
        s = raw.strip()
        if raw.startswith("# "):
            data["title"] = strip_markdown_inline(raw[2:].strip())
            idx += 1
            continue
        if raw.startswith("## 摘要"):
            flush_paragraph()
            current = "abstract"
            idx += 1
            continue
        if raw.startswith("## 第 "):
            flush_paragraph()
            current = "body"
            data["blocks"].append({"type": "h1", "text": strip_markdown_inline(raw[3:].strip())})
            idx += 1
            continue
        if raw.startswith("### ") and current == "body":
            flush_paragraph()
            data["blocks"].append({"type": "h2", "text": strip_markdown_inline(raw[4:].strip())})
            idx += 1
            continue
        if s.startswith("**关键词：**"):
            flush_paragraph()
            data["keywords"] = strip_markdown_inline(s.replace("**关键词：**", "").strip())
            idx += 1
            continue
        if not s:
            flush_paragraph()
            idx += 1
            continue
        if current == "body" and s.startswith("|"):
            table_lines: list[str] = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            if len(table_lines) >= 2 and re.search(r"^\|\s*[-:| ]+\|\s*$", table_lines[1]):
                flush_paragraph()
                rows = []
                for table_line in table_lines:
                    if re.search(r"^\|\s*[-:| ]+\|\s*$", table_line):
                        continue
                    cells = [strip_markdown_inline(cell.strip()) for cell in table_line.strip("|").split("|")]
                    rows.append(cells)
                if rows:
                    data["blocks"].append({"type": "table", "rows": rows})
                continue
            para_buf.extend(table_lines)
            continue
        if current == "body" and re.match(r"^\d+\.\s+", s):
            flush_paragraph()
            data["blocks"].append({"type": "list", "text": strip_markdown_inline(s)})
            idx += 1
            continue
        if current == "body" and s.startswith("- "):
            flush_paragraph()
            data["blocks"].append({"type": "bullet", "text": strip_markdown_inline(s[2:])})
            idx += 1
            continue
        para_buf.append(s)
        idx += 1

    flush_paragraph()
    return data


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_heading_style(style, *, font_name: str, size_pt: float, bold: bool, align) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_before = Pt(12 if size_pt >= 16 else 6)
    style.paragraph_format.space_after = Pt(12 if size_pt >= 16 else 6)
    style.paragraph_format.first_line_indent = Pt(0)
    style.paragraph_format.alignment = align


def ensure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(21)

    set_heading_style(doc.styles["Heading 1"], font_name="黑体", size_pt=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_heading_style(doc.styles["Heading 2"], font_name="宋体", size_pt=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_heading_style(doc.styles["Heading 3"], font_name="宋体", size_pt=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    for style_name in ["ThesisBody", "ThesisCaption", "ThesisAbstractTitle", "ThesisKeyword", "ThesisTableText"]:
        if style_name not in [s.name for s in doc.styles]:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    body_style = doc.styles["ThesisBody"]
    body_style.font.name = "宋体"
    body_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body_style.font.size = Pt(12)
    body_style.paragraph_format.line_spacing = 1.25
    body_style.paragraph_format.first_line_indent = Pt(21)
    body_style.paragraph_format.space_before = Pt(0)
    body_style.paragraph_format.space_after = Pt(0)

    cap_style = doc.styles["ThesisCaption"]
    cap_style.font.name = "宋体"
    cap_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap_style.font.size = Pt(10.5)
    cap_style.paragraph_format.line_spacing = 1.25
    cap_style.paragraph_format.space_before = Pt(6)
    cap_style.paragraph_format.space_after = Pt(6)
    cap_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    abs_style = doc.styles["ThesisAbstractTitle"]
    abs_style.font.name = "宋体"
    abs_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    abs_style.font.size = Pt(18)
    abs_style.font.bold = True
    abs_style.paragraph_format.line_spacing = 1.25
    abs_style.paragraph_format.space_before = Pt(12)
    abs_style.paragraph_format.space_after = Pt(12)
    abs_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    kw_style = doc.styles["ThesisKeyword"]
    kw_style.font.name = "宋体"
    kw_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    kw_style.font.size = Pt(12)
    kw_style.paragraph_format.line_spacing = 1.25
    kw_style.paragraph_format.space_before = Pt(6)
    kw_style.paragraph_format.space_after = Pt(0)

    table_style = doc.styles["ThesisTableText"]
    table_style.font.name = "宋体"
    table_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    table_style.font.size = Pt(10.5)
    table_style.paragraph_format.line_spacing = 1.15
    table_style.paragraph_format.space_before = Pt(0)
    table_style.paragraph_format.space_after = Pt(0)
    table_style.paragraph_format.first_line_indent = Pt(0)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="ThesisBody")
    set_paragraph_format(p, first_line_chars=2, line_spacing=1.25)
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)


def add_cover_page(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Cm(3.6))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本科毕业论文（设计）")
    set_run_font(run, "黑体", 22, bold=True)
    set_paragraph_format(p, before_pt=10, after_pt=24)

    table = doc.add_table(rows=9, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(3.8), Cm(1.0), Cm(11.0)]
    data = [
        ("中文题目", "：", TITLE_CN),
        ("英文题目", "：", TITLE_EN),
        ("姓名", "：", AUTHOR),
        ("学号", "：", STUDENT_ID),
        ("班级", "：", CLASS_NAME),
        ("专业", "：", MAJOR),
        ("学院", "：", COLLEGE),
        ("指导教师", "：", ADVISOR),
        ("完成时间", "：", f"{GRAD_YEAR}年4月"),
    ]

    for row_idx, row in enumerate(table.rows):
        row.height = Cm(1.1)
        for col_idx, cell in enumerate(row.cells):
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(para, line_spacing=1.15)
            para.style = doc.styles["ThesisTableText"]
            value = data[row_idx][col_idx]
            run = para.add_run(value)
            if col_idx == 0:
                set_run_font(run, "宋体", 12, bold=True)
            else:
                set_run_font(run, "宋体", 11)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SCHOOL)
    set_run_font(run, "宋体", 14, bold=True)


def add_placeholder_page(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(title)
    set_run_font(run, "黑体", 16, bold=True)
    add_body_paragraph(doc, body)


def add_abstract_pages(doc: Document, abstract: list[str], keywords: str) -> None:
    p = doc.add_paragraph(style="ThesisAbstractTitle")
    run = p.add_run("摘  要")
    set_run_font(run, "宋体", 18, bold=True)
    for para in abstract:
        add_body_paragraph(doc, para)
    p = doc.add_paragraph(style="ThesisKeyword")
    r1 = p.add_run("关键词：")
    set_run_font(r1, "黑体", 12, bold=True)
    r2 = p.add_run(keywords or "待补充")
    set_run_font(r2, "宋体", 12)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True
    set_paragraph_format(p, before_pt=12, after_pt=12)
    add_body_paragraph(doc, "English abstract to be finalized based on the Chinese abstract content.")
    p = doc.add_paragraph(style="ThesisKeyword")
    r1 = p.add_run("Key words: ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run("Smart weather service; multi-model forecast; upper-air analysis; historical climate; machine learning")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)


def add_toc_page(doc: Document) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    p = doc.add_paragraph(style="ThesisAbstractTitle")
    run = p.add_run("目  录")
    set_run_font(run, "宋体", 18, bold=True)
    toc_para = doc.add_paragraph()
    add_toc_field(toc_para)


def add_heading(doc: Document, text: str, level: int) -> None:
    style_name = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style_name)
    p.clear()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", 16, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        set_run_font(run, "宋体", 14, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    keep_paragraph_together(p, with_next=True)


def add_image_with_caption(doc: Document, image_path: Path, caption: str, width_cm: float,
                           *, page_break_before: bool = False, max_height_cm: float | None = None) -> None:
    if page_break_before:
        doc.add_page_break()
    if not image_path.exists():
        add_body_paragraph(doc, f"[插图缺失：{caption}，预期文件：{image_path.name}]")
        return

    target_width_cm = width_cm
    if max_height_cm is not None:
        image = Image.from_file(str(image_path))
        width_px = float(image.px_width)
        height_px = float(image.px_height)
        if width_px > 0 and height_px > 0:
            scaled_height_cm = width_cm * (height_px / width_px)
            if scaled_height_cm > max_height_cm:
                target_width_cm = max_height_cm * (width_px / height_px)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(target_width_cm))
    keep_paragraph_together(p, with_next=True)
    cp = doc.add_paragraph(style="ThesisCaption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(caption)
    set_run_font(run, "宋体", 10.5)
    keep_paragraph_together(cp, with_next=False)


def insert_assets_for_heading(doc: Document, heading_text: str) -> None:
    diagrams = DOCS / "diagrams" / "images"
    screenshots = DOCS / "images"
    mapping = {
        "4.1 系统总体架构": [
            (diagrams / "system_architecture.png", "图 4-1  智能气象信息服务系统总体架构图", 14.8, False),
            (diagrams / "data_flow.png", "图 4-2  系统核心数据流图", 13.2, False),
            (diagrams / "module_relationship.png", "图 4-3  系统主要模块关系图", 13.5, False),
            (screenshots / "home1.png", "图 4-4  首页导航页界面示意图", 14.8, False),
        ],
        "5.1 多模式预报模块设计与实现": [
            (diagrams / "advanced_forecast_flow.png", "图 5-1  多模式预报模块流程图", 11.0, False),
            (screenshots / "advanced-forecast1.png", "图 5-2  多模式预报页面示意图", 14.5, False),
        ],
        "5.2 探空分析模块设计与实现": [
            (diagrams / "upperair_flow.png", "图 5-3  探空分析模块流程图", 10.0, False),
            (screenshots / "upperair1.png", "图 5-4  探空分析页面示意图", 14.5, False),
        ],
        "5.3 历史气候分析模块设计与实现": [
            (diagrams / "history_ew_flow.png", "图 5-5  历史气候代表事件识别流程图", 10.0, False),
            (screenshots / "history.png", "图 5-6  历史气候年度页示意图", 14.5, False),
        ],
        "5.4 机器学习误差校正模块设计与实现": [
            (diagrams / "ml_correction_flow.png", "图 5-7  机器学习误差校正流程图", 9.6, False),
        ],
        "5.5 农业气象服务模块设计与实现": [
            (diagrams / "agro_service_flow.png", "图 5-8  农业气象服务模块流程图", 11.0, False),
            (screenshots / "agro-dashboard.png", "图 5-9  农业气象总览页示意图", 14.5, False),
        ],
    }
    for image_path, caption, width_cm, page_break_before in mapping.get(heading_text, []):
        max_height_cm = 16.5 if "flow" in image_path.name else 12.5 if image_path.suffix.lower() == ".png" and "flow" not in image_path.name and "home1" not in image_path.name else None
        add_image_with_caption(
            doc,
            image_path,
            caption,
            width_cm,
            page_break_before=page_break_before,
            max_height_cm=max_height_cm,
        )


def style_table(table, *, font_size: float = 10.5) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.style = "ThesisTableText"
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run_font(run, "宋体", font_size)


def set_table_column_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm[: len(row.cells)]):
            row.cells[idx].width = Cm(width)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    style_table(table, font_size=10)

    if col_count == 8:
        widths = [1.5, 1.8, 1.8, 1.8, 1.8, 1.7, 1.7, 2.1]
    elif col_count == 3:
        widths = [4.5, 4.0, 4.0]
    else:
        widths = [max(1.8, 15.0 / max(col_count, 1))] * col_count

    for row_values in rows:
        cells = table.add_row().cells
        for idx in range(col_count):
            cells[idx].text = row_values[idx] if idx < len(row_values) else ""

    set_table_column_widths(table, widths)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="ThesisCaption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", 10.5)


def add_metrics_tables(doc: Document, metrics: dict) -> None:
    reg = metrics["regression"]
    precip = metrics["precipitation"]["final_test"]

    add_table_caption(doc, "表 6-1  机器学习误差校正连续变量评估结果")
    table1 = doc.add_table(rows=1, cols=5)
    style_table(table1, font_size=10)
    headers = ["变量", "MAE（原始→校正）", "RMSE（原始→校正）", "Bias（原始→校正）", "改进幅度"]
    for cell, text in zip(table1.rows[0].cells, headers):
        cell.text = text
    set_table_column_widths(table1, [1.7, 3.2, 3.2, 3.2, 2.6])

    rows = [
        ("温度",
         f"{reg['temp']['final_test']['original_mae']:.3f}→{reg['temp']['final_test']['corrected_mae']:.3f}",
         f"{reg['temp']['final_test']['original_rmse']:.3f}→{reg['temp']['final_test']['corrected_rmse']:.3f}",
         f"{reg['temp']['final_test']['original_bias']:.3f}→{reg['temp']['final_test']['corrected_bias']:.3f}",
         f"{reg['temp']['final_test']['improvement_pct']:.1f}%"),
        ("湿度",
         f"{reg['rhum']['final_test']['original_mae']:.3f}→{reg['rhum']['final_test']['corrected_mae']:.3f}",
         f"{reg['rhum']['final_test']['original_rmse']:.3f}→{reg['rhum']['final_test']['corrected_rmse']:.3f}",
         f"{reg['rhum']['final_test']['original_bias']:.3f}→{reg['rhum']['final_test']['corrected_bias']:.3f}",
         f"{reg['rhum']['final_test']['improvement_pct']:.1f}%"),
        ("风速",
         f"{reg['wspd']['final_test']['original_mae']:.3f}→{reg['wspd']['final_test']['corrected_mae']:.3f}",
         f"{reg['wspd']['final_test']['original_rmse']:.3f}→{reg['wspd']['final_test']['corrected_rmse']:.3f}",
         f"{reg['wspd']['final_test']['original_bias']:.3f}→{reg['wspd']['final_test']['corrected_bias']:.3f}",
         f"{reg['wspd']['final_test']['improvement_pct']:.1f}%"),
    ]
    for row in rows:
        cells = table1.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text

    add_table_caption(doc, "表 6-2  机器学习误差校正降水评估结果")
    table2 = doc.add_table(rows=1, cols=7)
    style_table(table2, font_size=9.5)
    headers = ["Accuracy", "Precision", "Recall", "F1", "雨样本 MAE", "全样本 MAE", "说明"]
    for cell, text in zip(table2.rows[0].cells, headers):
        cell.text = text
    set_table_column_widths(table2, [1.7, 1.7, 1.6, 1.4, 1.9, 1.9, 3.6])
    row = table2.add_row().cells
    values = [
        f"{precip['classification_raw']['accuracy']:.3f}→{precip['classification_corrected']['accuracy']:.3f}",
        f"{precip['classification_raw']['precision']:.3f}→{precip['classification_corrected']['precision']:.3f}",
        f"{precip['classification_raw']['recall']:.3f}→{precip['classification_corrected']['recall']:.3f}",
        f"{precip['classification_raw']['f1']:.3f}→{precip['classification_corrected']['f1']:.3f}",
        f"{precip['rainy_mae_raw']:.3f}→{precip['rainy_mae_corrected']:.3f}",
        f"{precip['overall_mae_raw']:.3f}→{precip['overall_mae_corrected']:.3f}",
        "偏保守，雨量更准",
    ]
    for cell, text in zip(row, values):
        cell.text = text


def build_body(doc: Document, blocks: Iterable[dict], metrics: dict) -> None:
    current_h1 = ""
    for block in blocks:
        btype = block["type"]
        if btype == "h1":
            text = block["text"]
            if current_h1:
                doc.add_page_break()
            current_h1 = text
            add_heading(doc, text, 1)
        elif btype == "h2":
            text = block["text"]
            add_heading(doc, text, 2)
            insert_assets_for_heading(doc, text)
            if text == "6.4 模型评估结果分析":
                add_metrics_tables(doc, metrics)
        elif btype == "paragraph":
            text = block["text"]
            add_body_paragraph(doc, text)
        elif btype == "table":
            add_markdown_table(doc, block["rows"])
        elif btype == "list":
            text = block["text"]
            add_body_paragraph(doc, text)
        elif btype == "bullet":
            text = block["text"]
            add_body_paragraph(doc, f"• {text}")


def write_asset_manifest() -> None:
    manifest = FINAL_DIR / "asset_manifest.md"
    content = f"""# 正式论文素材清单

## 正文来源

- `docs\\thesis_draft.md`

## 模板来源

- `docs/final_paper/school_template_source.docx`
- 原始学校模板：`docs/学校模板.docx`

## 当前插图来源

### 结构图 / 流程图

- `docs/diagrams/images/system_architecture.png`
- `docs/diagrams/images/data_flow.png`
- `docs/diagrams/images/module_relationship.png`
- `docs/diagrams/images/advanced_forecast_flow.png`
- `docs/diagrams/images/upperair_flow.png`
- `docs/diagrams/images/history_ew_flow.png`
- `docs/diagrams/images/ml_correction_flow.png`
- `docs/diagrams/images/agro_service_flow.png`

### 系统截图

- `docs/images/home1.png`
- `docs/images/advanced-forecast1.png`
- `docs/images/upperair1.png`
- `docs/images/history.png`
- `docs/images/agro-dashboard.png`

## 待人工补充字段

- 学院
- 专业
- 班级
- 指导教师
- 最终英文摘要
- 最终参考文献条目

## 模板复刻状态

- 已改为双题目封面表格
- 已插入模板内嵌校徽
- 已切换为 Word Heading 标题体系
- 已拆分连续变量和降水评估表
"""
    manifest.write_text(content, encoding="utf-8")


def add_references_section(doc: Document) -> None:
    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Grinberg M. Flask Web Development: Developing Web Applications with Python[M]. 2nd ed. Sebastopol: O'Reilly Media, 2018.",
        "[2] Harris C R, Millman K J, van der Walt S J, et al. Array programming with NumPy[J]. Nature, 2020, 585(7825): 357-362.",
        "[3] McKinney W. Data Structures for Statistical Computing in Python[C]//Proceedings of the 9th Python in Science Conference. Austin: SciPy, 2010: 56-61.",
        "[4] Hunter J D. Matplotlib: A 2D graphics environment[J]. Computing in Science & Engineering, 2007, 9(3): 90-95.",
        "[5] May R M, Goebbert K H, Thielen J E, et al. MetPy: A Meteorological Python Library for Data Analysis and Visualization[J]. Bulletin of the American Meteorological Society, 2022, 103(10): E2273-E2284.",
        "[6] Munoz-Sabater J, Dutra E, Agusti-Panareda A, et al. ERA5-Land: A state-of-the-art global reanalysis dataset for land applications[J]. Earth System Science Data, 2021, 13(9): 4349-4383.",
        "[7] Zippenfenig P. Open-Meteo.com Weather API[CP/OL]. Zenodo, 2024[2026-04-14]. https://zenodo.org/records/14582479.",
        "[8] University of Wyoming Department of Atmospheric Science. University of Wyoming Atmospheric Science Radiosonde Archive[EB/OL]. [2026-04-14]. https://weather.uwyo.edu/upperair/sounding.shtml.",
        "[9] Glahn H R, Lowry D A. The use of model output statistics (MOS) in objective weather forecasting[J]. Journal of Applied Meteorology, 1972, 11(8): 1203-1211.",
        "[10] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "[11] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.",
        "[12] Ben Bouallegue Z, Clare M C A, Magnusson L, et al. The rise of data-driven weather forecasting: A first statistical assessment of machine learning-based weather forecasts in an operational-like context[J]. Bulletin of the American Meteorological Society, 2024, 105(6): E864-E883.",
        "[13] 唐卫亚, 等. 天气学分析基础[M]. 北京: 气象出版社, 2016.",
        "[14] 肖金香, 叶清, 吴仁烨. 农业气象学[M]. 第3版. 北京: 高等教育出版社, 2023.",
        "[15] 王春乙, 张继权, 霍治国, 等. 农业气象灾害风险评估研究进展与展望[J]. 气象学报, 2015(1): 1-19.",
        "[16] 中国气象局. 农业气象灾害风险区划技术导则: QX/T 527-2019[S]. 北京: 气象出版社, 2019.",
        "[17] 王春乙. 重大农业气象灾害研究进展[M]. 北京: 气象出版社, 2007.",
        "[18] 王春乙, 张雪芬, 赵艳霞. 农业气象灾害影响评估与风险评价[M]. 北京: 气象出版社, 2010.",
        "[19] 王春乙, 王石立, 霍治国. 近十余年来我国农业气象灾害研究进展[J]. 气象学报, 2005, 63(5): 659-668.",
        "[20] 李俊, 杜钧, 刘羽. 北京“7·21”特大暴雨不同集合预报方案的对比试验[J]. 气象学报, 2015, 73(1): 50-71.",
    ]

    for ref in refs:
        p = doc.add_paragraph(style="ThesisBody")
        p.paragraph_format.first_line_indent = Pt(-21)
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(ref)
        set_run_font(run, "宋体", 12)


def save_outputs(doc: Document) -> None:
    doc.save(str(OUTPUT_ASCII))


def update_docx_fields_via_word(path: Path) -> None:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return

    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(path))
        doc.Fields.Update()
        for toc in doc.TablesOfContents:
            toc.Update()
        doc.Save()
        doc.Close(False)
    except Exception as exc:
        print(f"Field update skipped for {path.name}: {exc}")
    finally:
        if word is not None:
            word.Quit()


def main() -> None:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    draft_text = DRAFT_PATH.read_text(encoding="utf-8")
    parsed = parse_thesis_markdown(draft_text)
    metrics = json.loads(METRICS_PATH.read_text(encoding="utf-8"))

    doc = Document(str(TEMPLATE_PATH)) if TEMPLATE_PATH.exists() else Document()
    clear_document_body(doc)
    doc.settings.odd_and_even_pages_header_footer = True
    ensure_styles(doc)

    cover_sec = doc.sections[0]
    configure_section(cover_sec, with_header=False, with_footer_page_number=False, page_fmt=None)
    add_cover_page(doc)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1], with_header=False, with_footer_page_number=False, page_fmt=None)
    add_placeholder_page(doc, "学术承诺与使用授权说明", "本页按学校要求保留，作者签名、指导教师签名与日期建议在最终提交前手写补充。")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1], with_header=True, with_footer_page_number=True, page_fmt="upperRoman", page_start=1)
    add_abstract_pages(doc, parsed["abstract"], parsed["keywords"])
    configure_section(doc.sections[-1], with_header=True, with_footer_page_number=True, page_fmt="upperRoman")

    add_toc_page(doc)
    configure_section(doc.sections[-1], with_header=True, with_footer_page_number=True, page_fmt="upperRoman")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(doc.sections[-1], with_header=True, with_footer_page_number=True, page_fmt="decimal", page_start=1)
    build_body(doc, parsed["blocks"], metrics)

    doc.add_page_break()
    add_references_section(doc)

    doc.add_page_break()
    add_heading(doc, "致谢", 1)
    add_body_paragraph(doc, "致谢部分待作者根据实际情况补充。")

    save_outputs(doc)
    update_docx_fields_via_word(OUTPUT_ASCII)
    if OUTPUT_CN.exists():
        update_docx_fields_via_word(OUTPUT_CN)
    write_asset_manifest()
    print(f"Generated: {OUTPUT_ASCII}")


if __name__ == "__main__":
    main()
