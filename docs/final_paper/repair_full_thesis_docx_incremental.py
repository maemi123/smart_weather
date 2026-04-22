from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


FINAL_DIR = Path(r"D:\pythonProject\smart_weather\docs\final_paper")
TARGET = FINAL_DIR / "毕业论文初稿完全版.docx"
BASELINE = FINAL_DIR / "full_thesis_before_incremental_20260422_160558.docx"


def backup_target() -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = FINAL_DIR / f"full_thesis_before_incremental_{ts}.docx"
    backup.write_bytes(TARGET.read_bytes())
    return backup


def restore_baseline() -> None:
    TARGET.write_bytes(BASELINE.read_bytes())


def set_run_font(run, size=10.5, bold=False) -> None:
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def format_paragraph(paragraph, style: str, text: str, *, align=None, first_line=True, bold=False) -> None:
    paragraph.clear()
    paragraph.style = style
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=12 if style.startswith("Heading") else 10.5, bold=bold)
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.25
    if style == "ThesisBody" and first_line:
        pf.first_line_indent = Pt(21)
    else:
        pf.first_line_indent = None


def insert_paragraph_after(paragraph, style: str, text: str, *, align=None, first_line=True, bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    format_paragraph(new_para, style, text, align=align, first_line=first_line, bold=bold)
    return new_para


def insert_table_after(paragraph, rows: int, cols: int):
    doc = paragraph._parent.part.document
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def set_cell(cell, text: str, *, bold=False, center=True) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def find_para(doc: Document, startswith: str):
    for para in doc.paragraphs:
        if para.text.strip().startswith(startswith):
            return para
    raise ValueError(f"Paragraph not found: {startswith}")


def add_27(doc: Document) -> None:
    anchor = find_para(doc, "在评估方法上，系统采用按 issue_time")
    p = insert_paragraph_after(anchor, "Heading 2", "2.7 机器学习误差校正相关算法", first_line=False)
    p = insert_paragraph_after(
        p,
        "ThesisBody",
        "为了更好地说明本系统中误差校正模块的建模思路，本文对 V1 与 V2 两版方案涉及的典型算法作简要说明。前者主要采用随机森林完成回归与分类任务，后者则尝试使用梯度提升树中的 CatBoost 方案开展对比实验。",
    )
    p = insert_paragraph_after(
        p,
        "ThesisBody",
        "（1）随机森林（Random Forest）。随机森林是一种基于决策树的集成学习方法，其核心思想是通过 Bootstrap 重采样生成多个训练子集，再分别训练多棵决策树，并在预测阶段对各棵树的结果进行集成。对于回归任务，模型通常输出各树结果的平均值；对于分类任务，则采用多数投票方式得到最终类别。该方法具有抗过拟合能力较强、能够处理非线性关系、对异常值和局部噪声不太敏感等优点，适合当前样本规模有限、业务特征较强的气象误差校正场景，因此本文 V1 模型采用了该算法。",
    )
    insert_paragraph_after(
        p,
        "ThesisBody",
        "（2）梯度提升树（CatBoost）。梯度提升树通过迭代方式不断训练新树来拟合前一轮模型的残差，从而逐步降低整体预测误差。CatBoost 是梯度提升树的一种高效工程实现，具有训练速度较快、对类别特征支持较好、泛化能力较强等特点。在本研究中，V2 实验模型尝试采用 CatBoost 对 ECMWF 误差校正任务进行重训，以验证在统一站点口径、统一 issue_time 规则和严格时间隔离评估条件下，boosting 类模型是否能够取得比 V1 更优的结果。",
    )


def add_54_block(doc: Document) -> None:
    anchor = find_para(doc, "当前系统不仅完成了模型训练")
    p = insert_paragraph_after(
        anchor,
        "ThesisBody",
        "为了更直观地说明两版模型的建模差异，本文将 V1 与 V2 在特征设计上的主要区别归纳如表 5-1 所示。",
    )
    cap = insert_paragraph_after(
        p,
        "ThesisCaption",
        "表 5-1  V1 与 V2 机器学习误差校正特征差异对比",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    table = insert_table_after(cap, 9, 3)
    rows = [
        ["特征类别", "V1", "V2"],
        ["预报值（温度、湿度、风速、降水）", "是", "是"],
        ["时间特征（小时、月份、星期）", "是", "是"],
        ["预报时效（lead_hours）", "是", "是"],
        ["循环编码（hour_sin/cos）", "是", "是"],
        ["季节（season）", "否", "是"],
        ["预报起报周期（issue_cycle，00Z/12Z）", "否", "是"],
        ["短期变化量（3h/6h forecast change）", "否", "是"],
        ["交互特征（lead × season）", "否", "是"],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell(table.cell(r, c), value, bold=(r == 0))
    insert_paragraph_after(
        cap,
        "ThesisBody",
        "V2 在 V1 基础上引入了季节、起报周期、短期变化量等特征，以更好地刻画误差的时序依赖性和季节性规律。与仅依赖基础预报值和通用时间变量相比，这类增强特征更强调同一模式误差会随季节、起报时次和预报时效共同变化的业务事实，因此更适合作为后续 boosting 实验方案的输入基础。",
    )


def add_64_block(doc: Document) -> None:
    anchor = find_para(doc, "降水变量的结果则表现出更强的复杂性")
    p = insert_paragraph_after(
        anchor,
        "ThesisBody",
        "为了进一步比较不同版本误差校正方案在严格时间隔离测试集上的表现，本文在原始预报、V1 与 V2 三种结果之间补充了统一对比，如表 6-3 所示。需要说明的是，该表对应的最终测试集按照 issue_time 留出，测试起报范围为 2026 年 4 月 4 日至 2026 年 4 月 10 日。",
    )
    cap = insert_paragraph_after(
        p,
        "ThesisCaption",
        "表 6-3  不同版本误差校正效果对比（严格时间隔离测试集）",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    table = insert_table_after(cap, 10, 5)
    rows = [
        ["气象要素", "指标", "Raw", "V1", "V2"],
        ["温度", "MAE", "2.355", "1.832", "2.220"],
        ["湿度", "MAE", "9.851", "9.085", "10.076"],
        ["风速", "MAE", "4.413", "2.613", "3.061"],
        ["降水分类", "Accuracy", "0.754", "0.743", "0.658"],
        ["降水分类", "Precision", "0.675", "0.677", "0.487"],
        ["降水分类", "Recall", "0.496", "0.426", "0.625"],
        ["降水分类", "F1", "0.572", "0.523", "0.548"],
        ["降水量（雨样本）", "MAE", "1.024", "0.586", "0.682"],
        ["降水量（全样本）", "MAE", "0.414", "0.251", "0.452"],
    ]
    best = {(1, 3), (2, 3), (3, 3), (8, 3), (9, 3)}
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell(table.cell(r, c), value, bold=(r == 0 or (r, c) in best))
    p2 = insert_paragraph_after(
        cap,
        "ThesisBody",
        "注：加粗为最优结果。V1 在温度、湿度、风速、降水量上均优于 Raw 和 V2。",
    )
    insert_paragraph_after(
        p2,
        "ThesisBody",
        "从结果可以看出，本轮 V2 实验并未优于 V1，这一现象是合理的。首先，在当前数据规模约 2 个月的条件下，随机森林这类 bagging 模型通常比 boosting 模型更稳定，对局部噪声和样本不均衡也更不敏感。其次，boosting 模型往往需要更多样本、更细的特征工程和更充分的参数调优，才能真正发挥优势，而本研究当前阶段更重视站点口径统一、issue_time 规则统一以及严格时间隔离评估流程的建立。因此，本文的核心结论并不是 V2 一定更强，而是说明在小规模业务样本条件下，随机森林仍然是更稳健的误差校正选择；同时，相比单纯更换模型，严格控制时间切分、避免目标时刻泄漏、保证训练与运行时口径一致，往往对结果可信度更为关键。",
    )


def main() -> None:
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    if not BASELINE.exists():
        raise FileNotFoundError(BASELINE)
    backup = backup_target()
    restore_baseline()
    doc = Document(str(TARGET))
    add_27(doc)
    add_54_block(doc)
    add_64_block(doc)
    doc.save(str(TARGET))
    print(f"TARGET={TARGET}")
    print(f"BACKUP={backup}")


if __name__ == "__main__":
    main()
